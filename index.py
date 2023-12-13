import flask
import os
import zipfile
from docx.api import Document
from flask import request, jsonify
from flask_s3 import *
from io import BytesIO
from werkzeug.utils import secure_filename
import boto3
from s3_credential import *
import pypandoc
from bs4 import BeautifulSoup
from datetime import datetime
import json
import uuid
import time



s3 = FlaskS3()


app = flask.Flask(__name__)
app.config["DEBUG"] = True
app.config['UPLOAD_FOLDER'] = os.path.join('word','media')
app.config['FLASKS3_BUCKET_NAME'] = os.environ.get("AWS_BUCKET_NAME")
app.config['AWS_ACCESS_KEY_ID'] = os.environ.get("AWS_ACCESS_KEY_ID")
app.config['AWS_SECRET_ACCESS_KEY'] = os.environ.get("AWS_SECRET_ACCESS_KEY")
app.config['FLASKS3_BUCKET_DOMAIN'] = 's3.ap-south-1.amazonaws.com'
s3.init_app(app)



s3_boto = boto3.client('s3')
s3_res = boto3.resource("s3")
buck = s3_res.Bucket(os.environ.get("AWS_BUCKET_NAME"))


def generate_unique_name(length=10):
    timestamp_got = datetime.now().strftime("%s")
    unique_name = f"img{timestamp_got}.png"
    return unique_name

def get_url(key):
    url = s3_boto.generate_presigned_url(
    ClientMethod='get_object',
    Params={
        'Bucket': os.environ.get("AWS_BUCKET_NAME"),
        'Key': key
    },
    ExpiresIn=9600)
    return url



@app.route("/upload-document",methods=["POST"])
def uploadDocument():
    document = request.files.get("file",False)
    if(document):
        memfile = BytesIO()
        document.save(memfile)
        document = Document(memfile)
        tables = document.tables
        z = zipfile.ZipFile(memfile)
        z.extractall()
        all_files = z.namelist()
        # print(all_files)
        images = filter(lambda x: x.startswith('/word/media/'), all_files)
        # return "yo"
        rels = {}
        real_name = {}
        for r in document.part.rels.values():          
            if isinstance(r._target, docx.parts.image.ImagePart):
                file_location = '/word/media/'+secure_filename(generate_unique_name())
                fbinary = open(f'word/media/{os.path.basename(r._target.partname)}',"rb")
                file_url_upload = os.path.join("/media/docimages",os.path.basename(file_location))
                s=buck.put_object(Body=fbinary.read(),Key=file_url_upload)
                rels[r.rId] = get_url(file_url_upload)
                # print(s.generate_presigned_url(expires_in=0))
                real_name[r.rId] = os.path.basename(r._target.partname)
        # Data will be a list of rows represented as dictionaries
        # containing each row's data.
        data = []
        keys = None
        topic_id = ''
        get_string = ""
        #print(dir(table.columns))
        for table in tables :
            tr = {}
            for i, row in enumerate(table.rows):
                tr[row.cells[0].text] = ''
                for paragraph in row.cells[1].paragraphs:
                    if(row.cells[0].text == 'Topic ID'):
                        topic_id = row.cells[1].text
                    for rId in rels:
                        if rId in paragraph._p.xml:
                            z.extract('word/media/'+real_name[rId],os.getcwd())
                            tr[row.cells[0].text]+=f'<img src="{rels[rId]}">'
                    if(row.cells[0].text == 'Problem Statement' or row.cells[0].text  == 'Correct Answer  Explanation') :
                        # print(paragraph.font.superscripipt)
                        print(paragraph._element.xml)
                        get_string+=paragraph._element.xml+"\n"
                        tr[row.cells[0].text]+='<p>'+paragraph.text+'</p>'
                       
                        # print(paragraph.style.font.superscript)
                    else :
                        tr[row.cells[0].text]+=paragraph.text
            data.append(tr)    
        allData = {}
        allData['document'] = {}
        allData['document']['Topic ID'] = topic_id
        allData['document']['questions'] = data
        with open("output.xml",'w') as file:
            file.write(get_string)            
        return jsonify(allData)



def arrangeData(data,variable,image_hash_data):
    if(data[1].findChild()):
        if(len(data[1].findAll("img"))):
            img_data = data[1].findAll("img")
            for i,img in enumerate(img_data):
                if(image_hash_data.get(os.path.basename(img["src"]))):    
                    object_url = "https://app.xxxx.com/{1}/{2}".format(
                        s3_boto.get_bucket_location(Bucket='xxxx-media')['LocationConstraint'],
                        'media/docimages',image_hash_data.get(os.path.basename(img["src"])))
                    img_data[i]["src"] = object_url
            variable[data[0].text] = str(data[1]).replace("<td>", "").replace("</td>", "")
        else:
            variable[data[0].text] = data[1].text
        
        # print(type(data[1]))
        # print(data[1])
    else:
        variable[data[0].text] = data[1].text
    return variable



def preprocessData(data,image_hash_data):
    tr_data = data.findAll("tr")
    result_recv = {}
    for i,tr in enumerate(tr_data):
        result = tr_data[i]["class"]
        if(result[0]!="header"):
            all_data = tr_data[i].findAll("td")
            # print(all_data[0],all_data[1])
            if(len(all_data)):
                arrangeData(all_data, result_recv,image_hash_data)
            else:
                all_data = tr.findAll("th")
                arrangeData(all_data, result_recv,image_hash_data)
    return result_recv



@app.route("/api/json",methods=["POST"])
def preprocessDocFunc():
    document = request.files.get("file",False)
    image_hash_data = {}
    errros_arr = []
    all_data = []
    topic_id = ""
    if(document):
        try:    
            document.save("static/predoc.docx")
            # document = Document(memfile)
            # tables = document.tables
            real_file_path = "static/predoc.docx"
            real_file_stream = open(real_file_path,"rb")
            z = zipfile.ZipFile(real_file_stream)
            z.extractall()
            all_files = z.namelist()
        
            # images_data = filter(lambda x:x.startwith("word/media"), all_files)
            for i in all_files:
            
                if(i.startswith("word/media")):
                    #unique_name = secure_filename(generate_unique_name())
                    unique_name = secure_filename(str(time.time())+uuid.uuid4().hex)
                    fbinary = open(os.path.join(os.getcwd(),f'word/media/{os.path.basename(i)}'),"rb")
                    file_url_upload = os.path.join("media/docimages",unique_name)
                    s=buck.put_object(Body=fbinary.read(),Key=file_url_upload)
                    image_hash_data[os.path.basename(i)] = unique_name
            html = pypandoc.convert_file(real_file_path, 'html',extra_args=['--webtex'])
            parser = BeautifulSoup(html,"html.parser").findAll("table")
            topic_id = (parser[0].find(text="Topic ID").findNext("th") if parser[0].find(text="Topic ID").findNext("th") else parser[0].find(text="Topic ID").findNext("td")).text
            all_data = [preprocessData(tdata,image_hash_data) for tdata in parser if preprocessData(tdata,image_hash_data)]
        except Exception as e:
            errros_arr.append(str(e))
        
        return {
            "document":{
                "Topic ID":topic_id,
                "questions":all_data
            },
            "errors":errros_arr
        }

@app.route("/api/html",methods=["POST"])
def htmlresponse():
    document = request.files.get("file",False)
    image_hash_data = {}
    errros_arr = []
    all_data = []
    topic_id = ""
    if(document):
        try:    
            document.save("static/predoc.docx")
            # document = Document(memfile)
            # tables = document.tables
            real_file_path = "static/predoc.docx"
            real_file_stream = open(real_file_path,"rb")
            z = zipfile.ZipFile(real_file_stream)
            z.extractall()
            all_files = z.namelist()
        
            # images_data = filter(lambda x:x.startwith("word/media"), all_files)
            for i in all_files:
            
                if(i.startswith("word/media")):
                    unique_name = secure_filename(str(time.time())+uuid.uuid4().hex)
                    #print(unique_name)
                    #exit()
                    fbinary = open(os.path.join(os.getcwd(),f'word/media/{os.path.basename(i)}'),"rb")
                    file_url_upload = os.path.join("media/docimages",unique_name)
                    s=buck.put_object(Body=fbinary.read(),Key=file_url_upload)
                    image_hash_data[os.path.basename(i)] = unique_name
                    time.sleep(1)
            html = pypandoc.convert(real_file_path,'html',extra_args=['--mathjax'])
            parser = BeautifulSoup(html,"html.parser").findAll("img")
            img_data = parser
            resp = str(html)
            for i,img in enumerate(img_data):
                
                if(image_hash_data.get(os.path.basename(img["src"]))):  
                    old_img = img_data[i]['src']
                    object_url = "https://app.xxxx.com/{1}/{2}".format(
                        s3_boto.get_bucket_location(Bucket='xxxx-media')['LocationConstraint'],
                        'media/docimages',image_hash_data.get(os.path.basename(img["src"])))
                    resp = resp.replace(old_img,object_url )
                    
        except Exception as e:
            errros_arr.append(str(e))
        
        return resp

if __name__ == '__main__':
  app.run(host='0.0.0.0', port=5000)

app.run()